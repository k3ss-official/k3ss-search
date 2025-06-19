import os
import sys
import platform
from pathlib import Path
from typing import List, Dict, Any
import mimetypes
import json
from datetime import datetime

# Document parsing imports
import PyPDF2
import pdfplumber
from docx import Document
import openpyxl
import csv

class FileDiscoveryService:
    """Service for discovering and searching files across different storage locations"""
    
    def __init__(self):
        self.supported_extensions = {
            '.txt', '.md', '.rtf', '.pdf', '.docx', '.doc', '.odt',
            '.xlsx', '.xls', '.csv', '.pptx', '.ppt', '.py', '.js',
            '.html', '.css', '.json', '.xml', '.yaml', '.yml'
        }
        
    def discover_storage_locations(self) -> List[Dict[str, Any]]:
        """Discover accessible storage locations on the system"""
        locations = []
        home_path = Path.home()
        
        # Always include home directory
        locations.append({
            'name': 'Home Directory',
            'path': str(home_path),
            'type': 'local',
            'accessible': True,
            'description': f'User home directory ({home_path})'
        })
        
        # Common cloud storage locations
        cloud_locations = self._detect_cloud_storage(home_path)
        locations.extend(cloud_locations)
        
        # External drives (Unix/Linux/Mac)
        if platform.system() in ['Darwin', 'Linux']:
            external_drives = self._detect_external_drives()
            locations.extend(external_drives)
            
        return locations
    
    def _detect_cloud_storage(self, home_path: Path) -> List[Dict[str, Any]]:
        """Detect common cloud storage locations with enhanced Google Drive detection"""
        cloud_locations = []
        
        # Enhanced Google Drive detection for macOS
        if platform.system() == 'Darwin':
            google_drive_paths = [
                # Standard Google Drive locations
                home_path / 'Google Drive',
                home_path / 'GoogleDrive',
                home_path / 'My Drive',
                # Google Drive File Stream locations
                home_path / 'Google Drive File Stream',
                home_path / 'GoogleDriveFileStream',
                # Google Drive for Desktop locations
                home_path / 'Google Drive for Desktop',
                # Check in /Volumes for mounted Google Drives
                Path('/Volumes') / 'GoogleDrive',
                Path('/Volumes') / 'Google Drive',
                Path('/Volumes') / 'My Drive'
            ]
            
            # Also check for Google Drive in /Volumes directory
            volumes_path = Path('/Volumes')
            if volumes_path.exists():
                try:
                    for volume in volumes_path.iterdir():
                        if volume.is_dir() and 'google' in volume.name.lower():
                            google_drive_paths.append(volume)
                except (PermissionError, OSError):
                    pass
        else:
            # Standard locations for other platforms
            google_drive_paths = [
                home_path / 'Google Drive',
                home_path / 'GoogleDrive',
                home_path / 'My Drive'
            ]
        
        # Check each potential Google Drive path
        for gd_path in google_drive_paths:
            if gd_path.exists() and gd_path.is_dir():
                try:
                    # Test if we can actually access the directory
                    list(gd_path.iterdir())
                    cloud_locations.append({
                        'name': f'Google Drive ({gd_path.name})',
                        'path': str(gd_path),
                        'type': 'cloud',
                        'accessible': True,
                        'description': f'Google Drive sync folder ({gd_path})'
                    })
                except (PermissionError, OSError):
                    # Add as inaccessible but detected
                    cloud_locations.append({
                        'name': f'Google Drive ({gd_path.name}) - Access Denied',
                        'path': str(gd_path),
                        'type': 'cloud',
                        'accessible': False,
                        'description': f'Google Drive detected but access denied ({gd_path})'
                    })
        
        # iCloud Drive (macOS)
        if platform.system() == 'Darwin':
            icloud_paths = [
                home_path / 'Library' / 'Mobile Documents' / 'com~apple~CloudDocs',
                home_path / 'iCloud Drive',
                home_path / 'iCloudDrive'
            ]
            
            for icloud_path in icloud_paths:
                if icloud_path.exists():
                    cloud_locations.append({
                        'name': 'iCloud Drive',
                        'path': str(icloud_path),
                        'type': 'cloud',
                        'accessible': True,
                        'description': f'iCloud Drive sync folder ({icloud_path})'
                    })
                    break
        
        # Dropbox
        dropbox_paths = [
            home_path / 'Dropbox',
            home_path / 'Dropbox (Personal)',
            home_path / 'Dropbox (Business)'
        ]
        
        for dropbox_path in dropbox_paths:
            if dropbox_path.exists():
                cloud_locations.append({
                    'name': f'Dropbox ({dropbox_path.name})',
                    'path': str(dropbox_path),
                    'type': 'cloud',
                    'accessible': True,
                    'description': f'Dropbox sync folder ({dropbox_path})'
                })
        
        # OneDrive
        onedrive_paths = [
            home_path / 'OneDrive',
            home_path / 'OneDrive - Personal',
            home_path / 'OneDrive - Business'
        ]
        
        for od_path in onedrive_paths:
            if od_path.exists():
                cloud_locations.append({
                    'name': f'OneDrive ({od_path.name})',
                    'path': str(od_path),
                    'type': 'cloud',
                    'accessible': True,
                    'description': f'OneDrive sync folder ({od_path})'
                })
                
        return cloud_locations
    
    def _detect_external_drives(self) -> List[Dict[str, Any]]:
        """Detect external drives and mounted volumes"""
        external_drives = []
        
        if platform.system() == 'Darwin':  # macOS
            volumes_path = Path('/Volumes')
            if volumes_path.exists():
                try:
                    for volume in volumes_path.iterdir():
                        if volume.is_dir() and volume.name not in ['Macintosh HD', 'Preboot', 'Recovery', 'VM', 'Data']:
                            # Skip Google Drive volumes as they're handled in cloud storage
                            if 'google' not in volume.name.lower():
                                external_drives.append({
                                    'name': f'External Drive ({volume.name})',
                                    'path': str(volume),
                                    'type': 'external',
                                    'accessible': True,
                                    'description': f'External volume ({volume})'
                                })
                except (PermissionError, OSError):
                    pass
        
        elif platform.system() == 'Linux':
            # Check common mount points
            mount_points = ['/media', '/mnt', '/run/media']
            for mount_point in mount_points:
                mount_path = Path(mount_point)
                if mount_path.exists():
                    try:
                        for user_dir in mount_path.iterdir():
                            if user_dir.is_dir():
                                for drive in user_dir.iterdir():
                                    if drive.is_dir():
                                        external_drives.append({
                                            'name': f'External Drive ({drive.name})',
                                            'path': str(drive),
                                            'type': 'external',
                                            'accessible': True,
                                            'description': f'External drive ({drive})'
                                        })
                    except (PermissionError, OSError):
                        pass
        
        return external_drives
    
    def search_files(self, search_paths: List[str], search_terms: List[str], 
                    search_content: bool = True, deep_search: bool = False) -> Dict[str, Any]:
        """Search for files containing specified terms with enhanced reporting"""
        results = []
        total_files_scanned = 0
        total_directories_scanned = 0
        skipped_files = 0
        
        for search_path in search_paths:
            path = Path(search_path)
            if not path.exists():
                continue
                
            for file_path in self._walk_directory(path):
                total_files_scanned += 1
                
                if file_path.suffix.lower() in self.supported_extensions:
                    file_info = self._analyze_file(file_path, search_terms, search_content, deep_search)
                    if file_info and file_info['matches']:
                        results.append(file_info)
                else:
                    skipped_files += 1
        
        return {
            'success': True,
            'results': results,
            'stats': {
                'total_files_scanned': total_files_scanned,
                'total_directories_scanned': total_directories_scanned,
                'matching_files': len(results),
                'skipped_files': skipped_files,
                'search_terms': search_terms,
                'deep_search_enabled': deep_search
            }
        }
    
    def _walk_directory(self, path: Path):
        """Recursively walk directory and yield file paths"""
        try:
            for item in path.rglob('*'):
                if item.is_file():
                    yield item
        except (PermissionError, OSError):
            # Skip directories we can't access
            pass
    
    def _analyze_file(self, file_path: Path, search_terms: List[str], 
                     search_content: bool, deep_search: bool = False) -> Dict[str, Any]:
        """Analyze a file for search terms and extract metadata"""
        try:
            file_info = {
                'path': str(file_path),
                'name': file_path.name,
                'size': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'type': self._get_file_type(file_path),
                'matches': [],
                'content_preview': '',
                'full_content': ''
            }
            
            # Check filename matches
            filename_lower = file_path.name.lower()
            for term in search_terms:
                if term.lower() in filename_lower:
                    file_info['matches'].append(term)
            
            # Check content matches if requested
            if search_content or deep_search:
                content = self._extract_text_content(file_path, deep_search)
                if content:
                    file_info['full_content'] = content
                    file_info['content_preview'] = content[:500] + '...' if len(content) > 500 else content
                    
                    content_lower = content.lower()
                    for term in search_terms:
                        if term.lower() in content_lower and term not in file_info['matches']:
                            file_info['matches'].append(term)
            
            return file_info if file_info['matches'] else None
            
        except (PermissionError, OSError, Exception):
            return None
    
    def _get_file_type(self, file_path: Path) -> str:
        """Get human-readable file type"""
        extension = file_path.suffix.lower()
        type_map = {
            '.pdf': 'PDF Document',
            '.docx': 'Word Document',
            '.doc': 'Word Document (Legacy)',
            '.txt': 'Text File',
            '.md': 'Markdown File',
            '.xlsx': 'Excel Spreadsheet',
            '.xls': 'Excel Spreadsheet (Legacy)',
            '.csv': 'CSV File',
            '.py': 'Python Script',
            '.js': 'JavaScript File',
            '.html': 'HTML File',
            '.css': 'CSS File',
            '.json': 'JSON File'
        }
        return type_map.get(extension, f'{extension.upper()} File')
    
    def _extract_text_content(self, file_path: Path, deep_search: bool = False) -> str:
        """Extract text content from various file types with optional deep search"""
        try:
            extension = file_path.suffix.lower()
            
            if extension == '.txt' or extension == '.md':
                return self._read_text_file(file_path)
            elif extension == '.pdf':
                return self._extract_pdf_text(file_path, deep_search)
            elif extension == '.docx':
                return self._extract_docx_text(file_path, deep_search)
            elif extension in ['.xlsx', '.xls']:
                return self._extract_excel_text(file_path, deep_search)
            elif extension == '.csv':
                return self._extract_csv_text(file_path)
            elif extension in ['.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml']:
                return self._read_text_file(file_path)
            
        except Exception:
            pass
        
        return ''
    
    def _read_text_file(self, file_path: Path) -> str:
        """Read plain text file with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        return ''
    
    def _extract_pdf_text(self, file_path: Path, deep_search: bool = False) -> str:
        """Extract text from PDF using pdfplumber with optional deep extraction"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ''
                max_pages = len(pdf.pages) if deep_search else min(10, len(pdf.pages))
                
                for i, page in enumerate(pdf.pages[:max_pages]):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                    except Exception:
                        # Skip problematic pages but continue
                        continue
                        
                return text
        except Exception:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''
                    max_pages = len(reader.pages) if deep_search else min(10, len(reader.pages))
                    
                    for page in reader.pages[:max_pages]:
                        try:
                            text += page.extract_text() + '\n'
                        except Exception:
                            continue
                    return text
            except Exception:
                return ''
    
    def _extract_docx_text(self, file_path: Path, deep_search: bool = False) -> str:
        """Extract text from Word document with optional deep extraction"""
        try:
            doc = Document(file_path)
            text = ''
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            
            # If deep search, also extract from tables, headers, footers
            if deep_search:
                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + '\t'
                        text += '\n'
                
                # Extract from headers and footers
                for section in doc.sections:
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            text += paragraph.text + '\n'
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            text += paragraph.text + '\n'
            
            return text
        except Exception:
            return ''
    
    def _extract_excel_text(self, file_path: Path, deep_search: bool = False) -> str:
        """Extract text from Excel file with optional deep extraction"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = ''
            
            sheets_to_process = workbook.sheetnames if deep_search else workbook.sheetnames[:3]
            
            for sheet_name in sheets_to_process:
                sheet = workbook[sheet_name]
                text += f'Sheet: {sheet_name}\n'
                
                max_rows = sheet.max_row if deep_search else min(100, sheet.max_row)
                
                for row in sheet.iter_rows(max_row=max_rows, values_only=True):
                    row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text += row_text + '\n'
                text += '\n'
            
            return text
        except Exception:
            return ''
    
    def _extract_csv_text(self, file_path: Path) -> str:
        """Extract text from CSV file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                text = ''
                for row in reader:
                    text += '\t'.join(row) + '\n'
                return text
        except Exception:
            return ''
    
    def format_for_llm(self, search_results: List[Dict[str, Any]], 
                      search_terms: List[str]) -> str:
        """Format search results for LLM consumption"""
        output = "=== DOCUMENT COLLECTION ===\n"
        output += f"Search Terms: {search_terms}\n"
        output += f"Total Files: {len(search_results)}\n"
        output += f"Collection Date: {datetime.now().isoformat()}\n\n"
        
        for i, result in enumerate(search_results, 1):
            output += f"--- FILE {i} ---\n"
            output += f"Path: {result['path']}\n"
            output += f"Name: {result['name']}\n"
            output += f"Type: {result['type']}\n"
            output += f"Size: {result['size']} bytes\n"
            output += f"Modified: {result['modified']}\n"
            output += f"Matches: {result['matches']}\n\n"
            
            if result['full_content']:
                output += "Content:\n"
                output += result['full_content']
                output += "\n\n"
            else:
                output += "Content: [Unable to extract text content]\n\n"
        
        output += "=== END COLLECTION ===\n"
        return output

